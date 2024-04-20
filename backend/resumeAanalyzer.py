import streamlit as st
import spacy
import re
import nltk
from nltk.corpus import stopwords
from nltk.stem import WordNetLemmatizer
from docx import Document
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity

nlp = spacy.load("en_core_web_lg")
skill_pattern_path = r"data\jz_complete_patterns.jsonl"
nltk.download(['stopwords', 'wordnet'])

class ResumeAnalyzer:

    @staticmethod
    def remove_extra_spaces(text):
        """        Remove extra spaces from the input text.

        It splits the input text into words, removes any extra spaces, and then joins the words back together with a single space.

        Args:
            text (str): The input text containing extra spaces.

        Returns:
            str: The cleaned text with extra spaces removed.
        """

        words = text.split()
        cleaned_text = ' '.join(words)
        return cleaned_text

    @staticmethod
    def clean_text(text):
        """        Clean the input text by removing special characters, URLs, and lemmatizing the words.

        Args:
            text (str): The input text to be cleaned.

        Returns:
            str: The cleaned text without special characters, URLs, and lemmatized words.
        """

        review = re.sub(
            r'(@[A-Za-z0-9]+)|([^0-9A-Za-z \t])|(https?://\S+)|^rt',
            " ",
            text,
        )
        review = review.lower()
        review = review.split()
        lm = WordNetLemmatizer()
        review = [
            lm.lemmatize(word)
            for word in review
            if not word in set(stopwords.words("english"))
        ]
        review = " ".join(review)
        return review

    def extract_contact_number(text):
        """        Extracts a contact number from the given text.

        This function searches for a valid contact number pattern within the input text and returns the first match found.

        Args:
            text (str): The input text from which the contact number is to be extracted.

        Returns:
            str or None: The extracted contact number if found, otherwise None.
        """

        pattern = r"\b(?:\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b"
        match = re.search(pattern, text)
        return match.group() if match else None

    def extract_email(text):
        """        Extracts an email address from the given text.

        This function searches for an email address pattern within the input text and returns the first match found.

        Args:
            text (str): The input text from which the email address is to be extracted.

        Returns:
            str or None: The extracted email address if found, otherwise None.
        """

        pattern = r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b"
        match = re.search(pattern, text)
        return match.group() if match else None

    def extract_skills(text, nlp):
        """        Extract skills from the input text using spaCy NLP model.

        Args:
            text (str): The input text from which skills need to be extracted.
            nlp (spacy.language.Language): The spaCy NLP model to be used for text processing.

        Returns:
            list: A list of skills extracted from the input text.
        """

        doc = nlp(text)
        return [ent.text for ent in doc.ents if ent.label_ == "SKILL"]
    
    def unique_skills(text):
        """        Return a list of unique skills from the input text.

        This function takes a string of skills as input and returns a list of unique skills found in the input text.

        Args:
            text (str): A string containing skills.

        Returns:
            list: A list of unique skills extracted from the input text.
        """

        return list(set(text))

    def extract_soft_skills(text,nlp):
        """        Extract soft skills from the input text using spaCy NLP model.

        This function takes the input text and a spaCy NLP model as input and returns a list of soft skills extracted from the text.

        Args:
            text (str): The input text from which soft skills are to be extracted.
            nlp (spacy.Language): The spaCy NLP model used for text processing.

        Returns:
            list: A list of soft skills extracted from the input text.
        """

        doc = nlp(text)
        return [ent.text for ent in doc.ents if ent.label_ == "SOFT_SKILL"]

    def extract_entities(text,nlp):
        """        Extract entities from the given text using the provided NLP model.

        This function takes the input text and an NLP model, processes the text to extract entities, and categorizes them into different labels.

        Args:
            text (str): The input text from which entities need to be extracted.
            nlp (spacy.language): The NLP model used for entity extraction.

        Returns:
            dict: A dictionary containing lists of entities categorized by labels.
        """

        entities = {label: [] for label in ["Job-Category", "ORG", "EDU", "GPE", "PK_ORG"]}
        doc = nlp(ResumeAnalyzer.clean_text(text))
        for ent in doc.ents:
            if ent.label_ in entities:
                entities[ent.label_].append(ent.text)
        return entities


    def extract_text_from_word_document(docx_file):
        """        Extracts text from a Word document and returns the full text.

        It extracts the header text and main content text from the Word document and concatenates them to form the full text.

        Args:
            docx_file (str): The file path of the Word document.

        Returns:
            str: The full text extracted from the Word document.

        Raises:
            Exception: If an error occurs during the extraction process.
        """

        try:
            doc = Document(docx_file)
            
            # Extract header text
            headers = [section.header for section in doc.sections]
            header_text = ""
            for header in headers:
                for paragraph in header.paragraphs:
                    header_text += paragraph.text + "\n"
            
            # Extract main content text
            paragraphs = [paragraph.text for paragraph in doc.paragraphs]
            main_content_text = ' '.join(paragraphs)
            
            # Concatenate header and main content text
            full_text = header_text + "\n" + main_content_text
            
            return full_text
        except Exception as e:
            print(f"Error: {e}")
            return None

    def job_matching_algorithm(resume, job_description):
        """        Calculate the match percentage between a resume and a job description.

        This function takes a resume and a job description as input, calculates the match percentage
        between them using cosine similarity, and returns the match percentage.

        Args:
            resume (str): The text of the resume.
            job_description (str): The text of the job description.

        Returns:
            float: The match percentage between the resume and the job description.

        Raises:
            Exception: If an error occurs during the job matching algorithm.
        """

        if resume is None or job_description is None:
            return None
        text = [resume, job_description]
        cv = CountVectorizer()
        count_matrix = cv.fit_transform(text)
        try:
            similarity_scores = cosine_similarity(count_matrix)
            match_percentage = similarity_scores[0][1] * 100
            match_percentage = round(match_percentage, 2)
            return match_percentage
        except Exception as e:
            print(f"Error in job matching algorithm: {e}")
            return None

    @staticmethod
    def job_matching_feedback(match_percentage):
        """        Determine the job matching feedback based on the match percentage.

        This function takes in the match percentage and returns a feedback based on the following criteria:
        - If match_percentage is greater than or equal to 90, it returns "Master".
        - If match_percentage is between 80 and 90 (exclusive), it returns "Expert".
        - If match_percentage is between 70 and 80 (exclusive), it returns "Excellent".
        - If match_percentage is between 50 and 70 (exclusive), it returns "Good".
        - If match_percentage is between 40 and 50 (exclusive), it returns "Barely Passable".
        - If match_percentage is less than 40, it returns "Needs Improvement".

        Args:
            match_percentage (int): The percentage match of the job candidate with the job requirements.

        Returns:
            str: The feedback based on the match percentage.
        """

        if match_percentage >= 90:
            return "Master"
        elif 80 <= match_percentage < 90:
            return "Expert"
        elif 70 <= match_percentage < 80:
            return "Excellent"
        elif 50 <= match_percentage < 70:
            return "Good"
        elif 40 <= match_percentage < 50:
            return "Barely Passable"
        else:
            return "Needs Improvement"


    def display_demo_progress_bar():
        """        Display a circular progress bar with a random number.

        This function generates an HTML code to display a circular progress bar with a random percentage value.
        The HTML code is then styled using CSS to center align the content.
        """

        # Display the circular progress bar with a random number
        demo_percentage = 75
        demo_progress_bar_html = f"""  
        <div class="result-container">
            <br>
            <br>
            <br>
            <br>
            <br>
            <br>
            <br>
            <br>
            <br>
            <br>
            <div class="heading">
            </div>
            <div class="percentage-container">
                <div class="percentage-circle"></div>
                <span>{demo_percentage}%</span>
            </div>
            <h1>Demo</h1>
        </div>
        """
        # Define CSS to center align the content
        css = """
        <style>
        .result-container {
            margin: auto;

            width: 0%;
            text-align: center;
            justify-content: center;
        }
        </style>
        """
        st.markdown(css, unsafe_allow_html=True)
        st.markdown(demo_progress_bar_html, unsafe_allow_html=True)
